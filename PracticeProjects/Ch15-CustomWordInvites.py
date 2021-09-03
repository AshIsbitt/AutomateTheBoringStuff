# Ch15 Project - Custom invitations as word documents

import docx

# Open guests.txt
guests = open('guests.txt', 'r')

# Create empty word doc
doc = docx.Document('ch15-Invites.docx')

# For each attendee in guest.txt
guestList = guests.readlines()

for i in range(len(guestList)):
	# Write out text for invite
	# Apply appearance
	doc.add_paragraph('It would be a pleasure to have the company of', style='Custom1')

	# Add name from file
	doc.add_paragraph(guestList[i], style='Custom2')

	doc.add_paragraph('at 11010 Memory Lane on the Evening of', style='Custom1')
	doc.add_paragraph('April 1st', style='Custom3')
	doc.add_paragraph("at 7 o'clock", style='Custom1')

	# Add page break
	doc.paragraphs[4].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

# Save doc and close .txt
doc.save('ch15-Invites.docx')
guests.close()
