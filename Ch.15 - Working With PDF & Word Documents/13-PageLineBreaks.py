# Page and line breaks

import docx

doc = docx.Document()

doc.add_paragraph('This is on the first page!')

# Page breaks are complicated
doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

doc.add_paragraph('This is on the second page')

doc.save('13-twopage.docx')