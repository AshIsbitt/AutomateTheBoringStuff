# Adding Headings

import docx

doc = docx.Document()

# First param is text, second param applies the specific heading with that integer
doc.add_heading('Heading 0', 0)
doc.add_heading('Heading 1', 1)
doc.add_heading('Heading 2', 2)
doc.add_heading('Heading 3', 3)
doc.add_heading('Heading 4', 4)

doc.save('12-headings.docx')
