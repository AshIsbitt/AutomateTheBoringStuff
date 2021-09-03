# Writing word docs

import docx

doc = docx.Document()
doc.add_paragraph('Hello, world!')

paraObj1 = doc.add_paragraph('This is a second paragraph. ')
paraObj2 = doc.add_paragraph('This is a third!')
paraObj1.add_run('This text is added to the second paragraph...')

# The second parameter here applies the Title style to the text of the first parameter
doc.add_paragraph('HelloHelloAnthony', 'Title')

doc.save('11-helloworld.docx')